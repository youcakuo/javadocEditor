#!/usr/bin/env python
# coding: utf-8

# In[1]:


# coding: utf-8

import win32com.client
from docx.api import Document
import logging
import os
import re
import sys
from git import Repo
import csv
import ctypes
import socket
import sqlite3
import time

workingspace = ''

def get_spec_from_path(path, txnId):
    logging.debug('search from ' + path + ' for ' + txnId)
    result1 = ''
    result2 = ''
    for root, dirs, files in os.walk(path):
        for file in files:
            if file.endswith('.docx') and txnId in file:
                if not '歷史' in root and not '~$' in file:
                    result1 = os.path.join(root, file)
            if file.endswith('.doc') and txnId in file:
                if not '歷史' in root and not '~$' in file:
                    result2 = os.path.join(root, file)
    return result1 or result2

def get_spec_file(txnId):
    local_path = '.'
    spec_path = "\\\\10.204.1.80\\tfb\\T-NBTS 專案_第二階段\\工作區\\個人工作區\\Shen\\javadoc\\javadoc_spec"
    all_spec_path_80='\\\\10.204.1.80\\tfb\\T-NBTS 專案_第二階段\\!IISI_FILE\\81第二階段-需求規格書'
    all_spec_path_89='\\\\172.16.240.89\\E_Disk\\!IISI_FILE\\81第二階段-需求規格書'
    result = ''
    search_path = [local_path, spec_path, all_spec_path_89, all_spec_path_80]
    logging.debug(search_path)
    for path in search_path:
        result = get_spec_from_path(path, txnId)
        if result:
            break
    logging.debug(result)
    return result

def convert_doc_to_docx(file):
    if file.endswith('.doc'):
        word = win32com.client.Dispatch("Word.application")
        doc = word.Documents.Open(file)
        docx_file = '{0}{1}'.format(file, 'x')
        doc.SaveAs(docx_file, 12)
        doc.Close()
        word.Quit()
        logging.debug(docx_file)
        return docx_file
    else:
        logging.debug('NA')
        return ''
        
def specParser(txnId):
    logging.info('search spec for ' + txnId)
    spec_data1 = []
    spec_data2 = []
    spec_data3 = ''
    spec_data = []
    file = ''
    try:
        file = get_spec_file(txnId)
        logging.debug('get_spec_file: ' + file)
    except:
        logging.exception("Exception occurred when try to find sepcification")
    if not file:
        logging.warning('no valid spec found')
    else:
        converted_docx_file = convert_doc_to_docx(file)
        logging.info('sepcification:' + file)
        if converted_docx_file:
            logging.info('convert doc extension to docx')
        document = Document(converted_docx_file) if converted_docx_file else Document(file)
        logging.info('spec open done')
        for table in document.tables:
            spec_data1 = []
            spec_data2 = []
            keys = None
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)
                if i == 0:
                    keys = tuple(text)
                    if keys == ('序號', '序號欄位名稱', 'I/O', '資料型態', '畫面元件', '格式化', '預設值', '必輸', '唯讀', '隱藏', '屬性及檢核'):
                        logging.debug('欄位屬性')
                        continue
                    elif keys == ('序號', '欄位名稱', '處理方式'):
                        logging.debug('[specParser]欄位檢核說明')
                        continue
                    elif keys[0] == '交易初始化處理':
                        logging.debug('交易初始化處理')
                        tmp = keys[1].replace('N/A','').rstrip()
                        if tmp:
                            spec_data3 = tmp
                        break
                    else:
                        break
                row_data = dict(zip(keys, text))
                if '屬性及檢核' in keys and row_data['屬性及檢核']:
                    row_data['屬性及檢核'] = row_data['屬性及檢核'].replace('\n\n',',').replace('\n',',').replace(',,',',').replace('；,',',')
                    spec_data1.append(row_data)
                elif '處理方式' in keys and row_data['處理方式']:
                    row_data['處理方式'] = row_data['處理方式'].replace('\n\n',',').replace('\n',',').replace(',,',',').replace('；,',',')
                    spec_data2.append(row_data)
            if spec_data1 or spec_data2:
                spec_data.append(spec_data1)
                spec_data.append(spec_data2)
        spec_data.append(spec_data3)
    return spec_data
            
def fileNameToTxnId(fileName):
    ftokens = re.split(r'[/]', fileName)
    return ftokens[len(ftokens)-1][2:10]

def getSpec1Content(spec):
    result = ''
    if not spec:
        return result
    size = len(spec)-1
    isMutiple = False
    if size > 2:
        isMutiple = True
        preTab = '     *   '
    else:
        preTab = '     * '
    for i in range(size):
        if i%2 == 1:
            continue
        rowdata = spec[i]
        if isMutiple:
            result = result + '     * 第'+str(int(i/2))+'畫面:' + '\n'
        for row in rowdata:
            if len(row['屬性及檢核']) > 20:
                result = result + preTab + row['序號欄位名稱'].replace('\n','') + ': ' + row['屬性及檢核'] + '\n'
    logging.debug(result)
    return result
                
def getSpec2Content(spec):
    result = ''
    if not spec:
        return result
    size = len(spec)-1
    isMutiple = False
    if size > 2:
        isMutiple = True
        preTab = '     *   '
    else:
        preTab = '     * '
    for i in range(size):
        if i%2 == 0:
            continue
        rowdata = spec[i]
        if isMutiple:
            result = result + '     * 第'+str(int(i/2))+'畫面:' + '\n'
        for row in rowdata:
            if len(row['處理方式']) > 20:
                result = result + preTab + row['欄位名稱'].replace('\n','') + ': ' + row['處理方式'] + '\n'
    logging.debug(result)
    return result
                
def getSpec3Content(spec):
    result = ''
    if not spec:
        return result
    if spec[len(spec) -1]:
        result =  '     * ' + spec[len(spec)-1]
        result = result.replace('\n','\n     * ') + '\n'
    logging.debug(result)
    return result
    
def read_user_defined_file(csv_path = ''):
    result = False
    if not csv_path:
        global workingspace
        csv_path = os.path.join(workingspace,'user.csv')
    logging.debug('try to load file: ' + csv_path)
    try:
        with open(csv_path, newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            result = True
    except Exception as e:
        logging.exception('error when try to open user.csv')
    return result
    
def get_user_defined_dict(func_name):
    global workingspace
    lines = []
    csv_path = os.path.join(workingspace,'user.csv')
    with open(csv_path, newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            op = row['relation']
            key = row['key']
            l3 = row['line3']
            l4 = row['line4']
            match = False
            if op == 'eq':
                if func_name == key:
                    match = True
            elif op == 'bg':
                if func_name.startswith(key):
                    match = True
            elif op == 'ed':
                if func_name.endswith(key):
                    match = True
            elif op == 'in':
                if key in func_name:
                    match = True
            else:
                logging.error('unknown relation ' + op + ' when matching ' + key)
                continue
            if match:
                logging.debug('match ' + func_name + ' with ' + key + ' for ' + op)
                if l3:
                    lines.append(l3)
                    if l4:
                        lines.append(l4)
            else:
                logging.debug('not match ' + func_name + ' with ' + key + ' for ' + op)
            if lines:
                break
    return lines
    
def find_modified_source_files():
    path = '.'
    java_files = [f for f in os.listdir(path) if f.endswith('.java')]
    repo_path = ''
    txn_dirs = ["c://iisi/infinity-developer/repos/infinity-application-tfbnbts-transactions","d://iisi/infinity-developer/repos/infinity-application-tfbnbts-transactions","c://iisi/develop/infinity-developer/repos/infinity-application-tfbnbts-transactions","d://iisi/develop/infinity-developer/repos/infinity-application-tfbnbts-transactions"]
    for txn_dir in txn_dirs:
        if os.path.isdir(txn_dir):
            repo_path = txn_dir
            break
    if not java_files:
        #os.chdir("c://iisi/infinity-developer/repos/infinity-application-tfbnbts-transactions")
        os.chdir(repo_path)
        repo = Repo('.')
        result = repo.git.status()
        tokens = re.split(r'[\n]', result)
        for token in tokens:
            if "modified:  " in token:
                filetokens = re.split(r'[ ]', token)
                fileToCheck = filetokens[len(filetokens) -1]
                if fileToCheck.endswith('.java'):
                    java_files.append(fileToCheck)  
    return java_files
    
def javadoc_template_for_sourcefile(javaFile):
    logging.info('process file: ' + javaFile)
    try:
        spec_data = specParser(fileNameToTxnId(javaFile))
    except Exception as e:
        spec_data = []
        logging.exception('Error when specParser')
        logging.debug(e)
    is_csv_found = read_user_defined_file()
    if not is_csv_found:
        logging.info('No user.csv found')
    else:
        logging.info('customization file user.csv found')
    newlines = []   
    scriptTag = ''
    funcName = ''
    parseStage = 0
    with open(javaFile, 'r', encoding='utf-8') as f:
        f_content = f.readlines()
        for i, line in enumerate(f_content):
            if '[method_name]' in str(line):
                scriptTag = ''
                funcName = ''
                parseStage = 1
                for j in range(20):
                    if '*' in str(f_content[i+j]):
                        continue
                    line1 = str(f_content[i+j])
                    if 'private' in line1 or 'public' in line1 or 'protected' in line1:
                        scriptTag = 'CommentScriptlet'
                    else:
                        for n in range(5):                            
                            line1 = str(f_content[i+j+n])
                            tokens = re.split( r'[@(]', line1 )
                            if tokens[1] == 'CommentScriptlet' or tokens[1] == 'RelationshipScriptlet':
                                scriptTag = tokens[1]
                                func_tokens = re.split(r'["]', tokens[len(tokens) -1])
                                funcName = func_tokens[len(func_tokens) - 2]
                            line1 = str(f_content[i+j+n+1])
                            if not '@' in line1:
                                if scriptTag == '':
                                    scriptTag = 'CommentScriptlet'
                                break
                    tokens = re.split(r'[(]', line1)
                    line1 = tokens[0]
                    tokens = re.split(r'[ ]', line1)
                    if not funcName:
                        funcName = tokens[len(tokens) -1]
                    break
                line = line.replace('[method_name]', '#' + scriptTag + ': ' + funcName)
                #newlines.append(line)
            elif '[override_name]' in str(line):
                scriptTag = ''
                funcName = ''
                parseStage = 1
                scriptTag = 'CommentScriptlet'
                for j in range(20):
                    if '*' in str(f_content[i+j]):
                        continue
                    line1 = str(f_content[i+j+1])
                    tokens = re.split(r'[(]', line1)
                    #print(tokens)
                    line1 = tokens[0]
                    tokens = re.split(r'[ ]', line1)
                    #print(tokens)
                    funcName = tokens[len(tokens)-1]
                    break
                line = line.replace('[override_name]', '#' + 'Method: ' + funcName)
                newlines.append(line)
                line = "     * #UsedByScriptlet: CrossValidation_Rule1\n"
                #newlines.append(line)
            elif parseStage == 2:
                if '@param f\n' in line:
                    line = line.replace('@param f', '@param f 流程Facade')
                elif '@param n\n' in line:
                    line = line.replace('@param n', '@param n 通知')
                elif '@param cs\n' in line:
                    line = line.replace('@param cs', '@param cs 來源交易內文')
                elif '@param ct\n' in line:
                    line = line.replace('@param ct', '@param ct 目的交易內文')
                elif '@param c\n' in line:
                    line = line.replace('@param c', '@param c 交易內文')
                elif '@throws Throwable\n' in line:
                    line = line.replace('@throws Throwable', '@throws Throwable 例外錯誤')
                elif '@return\n' in line:
                    if scriptTag == 'RelationshipScriptlet':
                        line = line.replace('@return', '@return true為要執行關聯模組，false為不執行') 
                    elif funcName == 'InputController_1_FinishInputCondition':
                        line = line.replace('@return', '@return true為交易結束、false為尚未結束')
                    elif funcName == 'defaultBeforeInputConditions':
                        line = line.replace('@return', '@return true為可開放輸入、false為不可開放輸入')
            elif parseStage == 1:
                if '[method_desc]' in line or '[override_desc]' in line:
                    parseStage = 2
                    if is_csv_found:
                        try:
                            userDefined = get_user_defined_dict(funcName)
                        except Exception as e:
                            userDefined = None
                            is_csv_found = False
                            logging.exception('Error when read user.csv')
                    else:
                        userDefined = None
                    if scriptTag == 'RelationshipScriptlet':
                        newlines.append('     * 關聯條件運算式\n')
                    elif userDefined:
                        logging.debug('match from user custimization')
                        if len(userDefined) > 0:
                            newlines.append('     * ' + userDefined[0] + '\n')
                        if len(userDefined) > 1:
                            l4 = userDefined[1]
                            if '[PI]' == l4 and getSpec3Content(spec_data):
                                line = getSpec3Content(spec_data)
                            elif '[FC]' == l4 and getSpec1Content(spec_data):
                                line = getSpec1Content(spec_data)
                            elif '[CV]' == l4 and getSpec2Content(spec_data):
                                line = getSpec2Content(spec_data)
                            elif l4:
                                line = '     * ' + l4 + '\n'
                    elif funcName == 'doCrossValidationWhenAction':
                        newlines.append('     * 按鈕點擊的檢核與處理\n')
                    elif funcName == 'ActionControl':
                        newlines.append('     * 按鈕點擊的檢核與處理\n')
                    elif funcName == 'doCrossValidationWhenFieldInputCompleted':
                        newlines.append('     * 欄位輸入完畢時的檢核與處理\n')
                        if getSpec2Content(spec_data):
                            line = getSpec2Content(spec_data)
                    elif funcName == 'doCrossValidationWhenConfirmed':
                        newlines.append('     * 交易確認執行的檢核與處理\n')
                    elif funcName == 'FieldControl':
                        newlines.append('     * 欄位輸入完畢時的檢核與處理\n')
                        if getSpec1Content(spec_data):
                            line = getSpec1Content(spec_data)
                    elif funcName == 'FIELD_INPUT':
                        newlines.append('     * 欄位輸入完畢時的檢核與處理\n')
                        if getSpec1Content(spec_data):
                            line = getSpec1Content(spec_data)
                    elif 'ClientBeforeSendCBR003' in funcName:
                        newlines.append('     * 交易執行前的處理 (央媒查扣)\n')
                    elif 'ClientAfterSendCBR003' in funcName:
                        newlines.append('     * 交易執行後的處理 (央媒查扣)\n')
                    elif 'ClientBeforeSendCBR004' in funcName:
                        newlines.append('     * 交易執行前的處理 (央媒查扣沖正)\n')
                    elif 'ClientAfterSendCBR004' in funcName:
                        newlines.append('     * 交易執行後的處理 (央媒查扣沖正)\n')
                    elif 'ClientBeforeSend' in funcName:
                        newlines.append('     * 交易執行前的處理\n')
                    elif 'ClientAfterSend' in funcName:
                        newlines.append('     * 交易執行後的處理\n')
                    elif 'CBS' in funcName:
                        newlines.append('     * 交易執行前的處理\n')
                    elif 'CAS' in funcName:
                        newlines.append('     * 交易執行後的處理\n')
                    elif 'PatternInitial' in funcName:
                        newlines.append('     * 交易初始化\n')
                        if getSpec3Content(spec_data):
                            line = getSpec3Content(spec_data)
                    elif funcName == 'SetComposeTelegram':
                        newlines.append('     * 組合電文執行前的處理\n')
                    elif funcName == 'prepareCombineTelegram':
                        newlines.append('     * 組合電文執行前的處理\n')
                    elif funcName == 'defaultBeforeInputConditions':
                        newlines.append('     * 輸入模式開啟前的檢核與處理\n')
                    elif 'InputController' in funcName:
                        newlines.append('     * 交易確認執行完畢後，判斷是否應結束\n')
                        line = '     * 依據transactionState判斷交易是否結束\n'
                    #funcName = ''
            elif '*/' in line:
                parseStage = 0
            newlines.append(line)

    with open(javaFile, 'w', encoding='utf-8') as f:
        for line in newlines:
            f.write(line)
        logging.info('finish file: ' + javaFile + '\n')

def para_Handler(args):
    setDebug = False
    for arg in args:
        if arg.startswith('-'):
            arg1 = arg[1:]
            if arg1 == 'D':
                #logging.getLogger().setLevel(logging.DEBUG)
                logging.basicConfig(level=logging.DEBUG, format='[%(lineno)d]%(levelname)s: [%(funcName)s]%(message)s')
                setDebug = True
        else:
            print('Error parameter format')
    if not setDebug:
        logging.basicConfig(level=logging.INFO, format='%(message)s')
    logging.debug(args)
        
def message_to_show(count):
    show_message = ''
    r_count = -1
    try:
        conn = sqlite3.connect('\\\\10.204.1.80\\tfb\\T-NBTS 專案_第二階段\\工作區\\個人工作區\\Shen\\javadoc\\javadoc2.db')
        c = conn.cursor()
        cursor = c.execute("SELECT MESSAGE from UPDATEMESSAGE WHERE COUNT >= {}".format(count))
        for i, row in enumerate(cursor):
            if row[0]:
                show_message += '- ' + row[0] + '\n'
            r_count = i
    except Exception as e:
        logging.info(e)
    finally:
        conn.close() 
    if not show_message:
        show_message = 'no update information'
    if r_count >= 0:
        return [count+r_count+1, show_message]
    else:
        return None
    
def update_message():
    ipaddr = 'unknown IP address'
    current_ver = 0
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ipaddr = s.getsockname()[0]
    except OSError as e:
        logging.info(e)
    finally:
        s.close()
    userid = os.environ['COMPUTERNAME']
    found_record = False
    logging.info('userid: ' + userid + ' at ' + ipaddr)
    try:
        conn = sqlite3.connect('\\\\10.204.1.80\\tfb\\T-NBTS 專案_第二階段\\工作區\\個人工作區\\Shen\\javadoc\\javadoc2.db')
        c = conn.cursor()
        c.execute('CREATE TABLE RECORD(IP CHAR(20) PRIMARY KEY NOT NULL,ID CHAR(50), VERSION int)')      
        conn.commit()
        logging.debug('create table')
    except Exception as e:
        logging.info(e)
    finally:
        conn.close()
    time.sleep(0.5)  
    try:
        conn = sqlite3.connect('\\\\10.204.1.80\\tfb\\T-NBTS 專案_第二階段\\工作區\\個人工作區\\Shen\\javadoc\\javadoc2.db')
        c = conn.cursor()
        cursor = c.execute("SELECT ID, VERSION from RECORD WHERE ID = '{}'".format(userid))
        data=cursor.fetchall()
        print(data)
        print(type(data))
        if len(data) > 0:
            found_record = True
            current_ver = data[0][1]
            logging.debug('record found')
        else:
            logging.debug('no record found')
    except Exception as e:
        logging.info(e)
    finally:
        conn.close()    
    time.sleep(0.5)    
    mts = message_to_show(current_ver)
    if mts:
        select = ctypes.windll.user32.MessageBoxW(0, mts[1], "更新資訊", 1)
        if select == 1:
            if not found_record:
                #print("INSERT INTO RECORD (IP,ID, VERSION) VALUES ('{}', '{}', '{}')".format(ipaddr, userid, mts[0]))
                try:
                    conn = sqlite3.connect('\\\\10.204.1.80\\tfb\\T-NBTS 專案_第二階段\\工作區\\個人工作區\\Shen\\javadoc\\javadoc2.db')
                    c = conn.cursor()
                    c.execute("INSERT INTO RECORD (IP,ID,VERSION) VALUES ('{}', '{}', '{}')".format(ipaddr, userid, mts[0]));
                    conn.commit()
                    logging.info('insert to db successful')
                except Exception as e:
                    logging.info(e)
                finally:
                    conn.close()
            else:
                #print("UPDATE RECORD SET VERSION = {} WHERE ID = '{}'".format(mts[0], userid))
                try:
                    conn = sqlite3.connect('\\\\10.204.1.80\\tfb\\T-NBTS 專案_第二階段\\工作區\\個人工作區\\Shen\\javadoc\\javadoc2.db')
                    c = conn.cursor()
                    c.execute("UPDATE RECORD SET VERSION = {} WHERE ID = '{}'".format(mts[0], userid));
                    conn.commit()
                    logging.info('update version successful')
                except Exception as e:
                    logging.info(e)
                finally:
                    conn.close()

def main():
    #logging.basicConfig(level=logging.DEBUG, format='[%(lineno)d]%(levelname)s: [%(funcName)s]%(message)s')
    #logging.getLogger().setLevel(logging.DEBUG)
    para_Handler(sys.argv[1:])
    update_message()
    global workingspace
    workingspace = os.getcwd()
    java_files = find_modified_source_files()
    print('modified files:')
    for mfiles in java_files:
        print('  ' + mfiles)
    print('----------------------------------------------------------------------------')
    for javaFile in java_files:
        javadoc_template_for_sourcefile(javaFile)
                
if __name__ == '__main__':
    try:
        main()
    finally:
        print('press Enter to continue...')
        input()

